# app/report/word_builder.py
from __future__ import annotations

import copy
import logging
import math
import os
import tempfile
from dataclasses import asdict, is_dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Sequence, Tuple

from docx import Document as _DocxDocument
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Один код сценария (как в POUO.code / SCENARIOS) → один файл шаблона в app/report/templates/
WORD_TEMPLATE_BY_SCENARIO: Dict[str, str] = {
    "POUO1": "templatePOUO1.docx",
    "POUO2": "templatePOUO2.docx",
    "POUO3": "templatePOUO3.docx",
    "POUO4": "templatePOUO4.docx",
    "POUO5": "templatePOUO5.docx",
    "POUO6": "templatePOUO6.docx",
}

# Титульный лист (при необходимости замените на реквизиты своего ВУЗа)
TITLE_PAGE_UNIVERSITY = (
    "Федеральное государственное автономное образовательное учреждение высшего образования"
)
TITLE_PAGE_UNIVERSITY_SHORT = "«Балтийский федеральный университет имени Иммануила Канта»"
TITLE_PAGE_FACULTY = "Институт физико-математических наук и информационных технологий"
TITLE_PAGE_DEPARTMENT = "Кафедра прикладной математики и информатики"
TITLE_PAGE_AUTHOR = "Крикун Александр Владиславович"
TITLE_PAGE_GROUP = "4ПМиИ"


def default_word_templates_dir() -> Path:
    return Path(__file__).resolve().parent / "templates"


def _one_template_code_from_project(project: Any) -> str:
    """
    Код сценария для выбора **одного** файла шаблона.

    Нужен только для ``resolve_word_template_path`` / ``render_report_for_project``.
    Смешанный проект → ошибка с подсказкой использовать ``render_full_report``.
    """
    pouos = getattr(project, "pouos", None) or []
    if not pouos:
        raise ValueError("В проекте нет ПОУО — нечего включать в отчёт.")
    codes = {p.code for p in pouos}
    if len(codes) != 1:
        raise ValueError(
            "Экспорт в один Word-файл без сборки возможен только при одном типе ПОУО. "
            f"Сейчас в проекте: {', '.join(sorted(codes))}. "
            "Для нескольких сценариев используйте полный отчёт (render_full_report)."
        )
    return next(iter(codes))


def word_template_filename_for_scenario(scenario_code: str) -> str:
    fn = WORD_TEMPLATE_BY_SCENARIO.get(scenario_code)
    if not fn:
        known = ", ".join(sorted(WORD_TEMPLATE_BY_SCENARIO))
        raise ValueError(
            f"Для сценария «{scenario_code}» не задан шаблон Word. "
            f"Известные сценарии: {known}."
        )
    return fn


def resolve_word_template_path(
    project: Any,
    templates_dir: str | Path | None = None,
) -> Tuple[Path, str]:
    """
    Возвращает (полный путь к файлу, имя файла) для **одного** типа ПОУО в ``project``.

    При нескольких разных ``code`` — ``ValueError`` (нужен :func:`render_full_report`).
    Нет файла — ``FileNotFoundError`` с указанием сценария и пути.
    """
    base = Path(templates_dir) if templates_dir else default_word_templates_dir()
    code = _one_template_code_from_project(project)
    filename = word_template_filename_for_scenario(code)
    path = base / filename
    if not path.is_file():
        raise FileNotFoundError(
            f"Не найден шаблон для сценария {code}: ожидался файл\n  {path}"
        )
    return path, filename


def word_template_debug_line(project: Any, templates_dir: str | Path | None = None) -> str:
    """Строка для отладки: шаблоны по каждому коду ПОУО (порядок как в project.pouos)."""
    base = Path(templates_dir) if templates_dir else default_word_templates_dir()
    _, order = group_pouos_by_code(project)
    if not order:
        return "Шаблон Word: в проекте нет ПОУО."
    lines: List[str] = ["Шаблоны Word (порядок первого появления кода в проекте):"]
    for code in order:
        try:
            filename = word_template_filename_for_scenario(code)
        except ValueError as e:
            lines.append(f"  • {code}: {e}")
            continue
        path = base / filename
        if path.is_file():
            lines.append(f"  • {code} → {filename} (файл найден)")
        else:
            lines.append(f"  • {code} → {filename} — нет файла:\n    {path}")
    return "\n".join(lines)

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Pt

from app.core.fuels import get_fuel

logger = logging.getLogger(__name__)


def _to_dict(obj: Any) -> Any:
    """
    Рекурсивно приводит dataclass / объект к dict / list.

    Зачем это нужно:
    - часть данных в проекте может быть dataclass;
    - часть может быть обычным dict;
    - для docxtpl удобнее работать уже с обычными словарями и списками.
    """
    if obj is None:
        return None
    if is_dataclass(obj):
        return asdict(obj)
    if isinstance(obj, dict):
        return {k: _to_dict(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_to_dict(x) for x in obj]
    if hasattr(obj, "__dict__"):
        return {k: _to_dict(v) for k, v in vars(obj).items()}
    return obj


def _round_if_number(v: Any, ndigits: int = 2) -> Any:
    """
    Если значение число — округляем.
    Если нет — возвращаем как есть.
    """
    if isinstance(v, (int, float)):
        return round(v, ndigits)
    return v


def _pretty_value(v: Any, ndigits: int = 2) -> str:
    """
    Красивое строковое представление значения для шаблона Word.
    """
    if v is None:
        return "не найдено"
    if isinstance(v, (int, float)):
        return str(round(v, ndigits))
    return str(v)


def _wind_zones_from_m_dot(m_dot: Any) -> Dict[str, Any]:
    """Эмпирические L и r₀ для облака ГВС (как в методике ПОУО2): L = 25√(M/W), r₀ = 12,5√(M/W)."""
    try:
        m = float(m_dot)
    except (TypeError, ValueError):
        m = 0.0
    if m <= 0:
        return {
            "L_wind1_m": None, "L_wind3_m": None,
            "r0_wind1_m": None, "r0_wind3_m": None,
        }

    def L(W: float) -> float:
        return round(25.0 * math.sqrt(m / W), 1)

    def r0(W: float) -> float:
        return round(12.5 * math.sqrt(m / W), 1)

    return {
        "L_wind1_m": L(1.0), "L_wind3_m": L(3.0),
        "r0_wind1_m": r0(1.0), "r0_wind3_m": r0(3.0),
    }


def _interp_q_kw_m2(table: Any, r_target: Any) -> float | None:
    """Линейная интерполяция q(r) по таблице факела."""
    if not table or r_target is None:
        return None
    try:
        rt = float(r_target)
    except (TypeError, ValueError):
        return None
    rows = []
    for r in table:
        try:
            rm = float(r.get("r_m"))
            qv = float(r.get("q_kw_m2"))
        except (TypeError, ValueError):
            continue
        rows.append((rm, qv))
    if not rows:
        return None
    rows.sort(key=lambda x: x[0])
    if rt <= rows[0][0]:
        return round(rows[0][1], 4)
    if rt >= rows[-1][0]:
        return round(rows[-1][1], 4)
    for i in range(len(rows) - 1):
        r0, q0 = rows[i]
        r1, q1 = rows[i + 1]
        if r0 <= rt <= r1:
            if r1 == r0:
                return round(q0, 4)
            t = (rt - r0) / (r1 - r0)
            return round(q0 + t * (q1 - q0), 4)
    return None


def _pretty_dict(d: Dict[str, Any], ndigits: int = 2) -> List[Dict[str, str]]:
    """
    Преобразует словарь в список словарей вида:
        [{"name": ..., "value": ...}, ...]

    Это удобно для циклов в docxtpl.
    """
    out = []
    for k, v in (d or {}).items():
        out.append({
            "name": str(k),
            "value": _pretty_value(v, ndigits),
        })
    return out


def _pretty_building_zones(d: Dict[str, Any]) -> List[Dict[str, str]]:
    """
    Подготовка зон повреждений зданий для Word.

    Ожидается, что зоны могут приходить как:
        "A": [r1, r2]
        "B": [r1, r2]
        ...

    Важная правка:
    если зона фактически отсутствует и расчёт дал [0.0, 0.0],
    то вместо "0.0–0.0 м" выводим "-".
    """
    out = []

    for k, v in (d or {}).items():
        if isinstance(v, (list, tuple)) and len(v) == 2:
            r1, r2 = v

            # Если границы зоны не определены
            if r1 is None or r2 is None:
                txt = "-"
            else:
                try:
                    r1f = float(r1)
                    r2f = float(r2)

                    # Если зона "нулевая", показываем прочерк
                    if r1f == 0.0 and r2f == 0.0:
                        txt = "-"
                    else:
                        r1s = _pretty_value(r1f)
                        r2s = _pretty_value(r2f)
                        txt = f"{r1s}–{r2s} м"
                except (TypeError, ValueError):
                    txt = _pretty_value(v)
        else:
            txt = _pretty_value(v)

        out.append({
            "name": str(k),
            "value": txt,
        })

    return out


def _pad_building_zones(zones: List[Dict[str, str]], target_len: int = 5) -> List[Dict[str, str]]:
    """
    Шаблон жёстко обращается к зонам так:
        p.tvs.zones_buildings[0]
        ...
        p.tvs.zones_buildings[4]

    Поэтому, если расчёт вернул меньше 5 зон,
    добиваем список заглушками, чтобы шаблон не падал.
    """
    out = list(zones or [])
    while len(out) < target_len:
        out.append({"name": "", "value": "-"})
    return out


def _safe_inline_image(doc: DocxTemplate, path: str, width_mm: int = 150):
    """
    Возвращает InlineImage, если файл существует.
    Если файла нет — возвращает None.

    Это позволяет в шаблоне писать:
        {% if p.tvs_dp_chart_img %}
        {{ p.tvs_dp_chart_img }}
        {% endif %}
    """
    if path and os.path.exists(path):
        return InlineImage(doc, path, width=Mm(width_mm))
    return None


class DotDict(dict):
    """
    Словарь с доступом по атрибуту для Jinja/docxtpl (p.inputs.lpg.M_kg_kmol).

    Обычный dict в Jinja не даёт .lpg — только ['lpg'], отсюда была ошибка
    «'dict object' has no attribute 'lpg'».

    Нет ключа — возвращается пустой DotDict (цепочка не падает); печать пустого — "".
    """

    def __getattr__(self, key: str) -> Any:
        if key.startswith("_"):
            raise AttributeError(key)
        try:
            return self[key]
        except KeyError:
            return DotDict()

    def __str__(self) -> str:
        return "" if len(self) == 0 else super().__str__()

    def __repr__(self) -> str:
        return "" if len(self) == 0 else super().__repr__()

    def __format__(self, format_spec: str) -> str:
        return "" if len(self) == 0 else str(self)


def _dot_for_jinja(obj: Any) -> Any:
    """
    Рекурсивно оборачивает dict → DotDict для шаблона; скаляры и InlineImage не трогаем.
    """
    if obj is None:
        return None
    if isinstance(obj, (str, int, float, bool)):
        return obj
    if isinstance(obj, InlineImage):
        return obj
    if isinstance(obj, DotDict):
        return obj
    if isinstance(obj, dict):
        return DotDict({k: _dot_for_jinja(v) for k, v in obj.items()})
    if isinstance(obj, list):
        return [_dot_for_jinja(x) for x in obj]
    if isinstance(obj, tuple):
        return tuple(_dot_for_jinja(x) for x in obj)
    return obj


def _build_release_block(results: Dict[str, Any]) -> Dict[str, Any]:
    """
    Подготовка блока release для шаблона Word.

    Здесь мы НЕ пересчитываем физику заново.
    Мы только:
    - берём уже рассчитанные значения из results["release"];
    - аккуратно округляем;
    - подготавливаем их под шаблон.
    """
    rel = (results or {}).get("release", {}) or {}

    return {
        "accident_pipe": rel.get("accident_pipe", ""),
        "P_up_kpa": _round_if_number(rel.get("P_up_kpa")),
        "P2_kpa": _round_if_number(rel.get("P2_kpa", rel.get("P_up_kpa"))),
        "d_hole_mm": _round_if_number(rel.get("d_hole_mm")),
        "d_m": _round_if_number(rel.get("d_m"), 4),
        "t_shutoff_s": _round_if_number(rel.get("t_shutoff_s")),

        "F_m2": _round_if_number(rel.get("F_m2"), 6),
        "v_g_m3_kg": _round_if_number(rel.get("v_g_m3_kg"), 6),
        "m_dot_kg_s": _round_if_number(rel.get("m_dot_kg_s", rel.get("G_kg_s")), 4),
        "M1T_kg": _round_if_number(rel.get("M1T_kg")),
        "sum_r2L_m3": _round_if_number(rel.get("sum_r2L_m3"), 6),
        "V2T_m3": _round_if_number(rel.get("V2T_m3")),
        "M2T_kg": _round_if_number(rel.get("M2T_kg")),
        "Mg_kg": _round_if_number(rel.get("Mg_kg")),
        "M_total_kg": _round_if_number(rel.get("M_total_kg", rel.get("Mg_kg"))),
        "m_cloud_kg": _round_if_number(rel.get("m_cloud_kg")),

        "Eud_J_kg": _round_if_number(rel.get("Eud_J_kg")),
        "E_concentration_correction": _round_if_number(rel.get("E_concentration_correction"), 6),
        "E_J": _round_if_number(rel.get("E_J")),
        "Z": _round_if_number(rel.get("Z"), 4),
        "rho_n_kg_m3": _round_if_number(rel.get("rho_n_kg_m3"), 4),
        "R0_J_kgK": _round_if_number(rel.get("R0_J_kgK"), 4),
        "T_K": _round_if_number(rel.get("T_K"), 2),

        # Поля, которые ждёт шаблон
        "L_wind1_m": _round_if_number(rel.get("L_wind1_m")),
        "L_wind3_m": _round_if_number(rel.get("L_wind3_m")),
        "r0_wind1_m": _round_if_number(rel.get("r0_wind1_m")),
        "r0_wind3_m": _round_if_number(rel.get("r0_wind3_m")),

        "skip_reason": rel.get("skip_reason"),
    }


def _jetfire_table_rows(raw_table: Any) -> List[Dict[str, Any]]:
    out = []
    for row in (raw_table or []):
        out.append({
            "r_m": _round_if_number(row.get("r_m")),
            "tau": _round_if_number(row.get("tau"), 6),
            "Fq": _round_if_number(row.get("Fq"), 6),
            "q_kw_m2": _round_if_number(row.get("q_kw_m2"), 4),
        })
    return out


def _jetfire_zones_rows(raw_zones: Any) -> List[Dict[str, str]]:
    zones = []
    for z in (raw_zones or []):
        zones.append({
            "q_thr_kw_m2": _round_if_number(z.get("q_thr_kw_m2"), 2),
            "r_m": _pretty_value(z.get("r_m")),
        })
    return zones


def _build_jetfire_block(results: Dict[str, Any]) -> Dict[str, Any]:
    """
    Подготовка блока факельного горения.
    Для СУГ (tank_park) в results["jet_fire"] может быть вложенный блок peak.
    """
    jf = (results or {}).get("jet_fire", {}) or {}
    params = jf.get("params", {}) or {}

    table = _jetfire_table_rows(jf.get("table"))
    zones = _jetfire_zones_rows(jf.get("zones"))

    out: Dict[str, Any] = {
        "params": {
            "M_kg_s": _round_if_number(params.get("M_kg_s"), 4),
            "LF_m": _round_if_number(params.get("LF_m"), 2),
            "DF_m": _round_if_number(params.get("DF_m"), 2),
            "Ef_kw_m2": _round_if_number(params.get("Ef_kw_m2"), 2),
        },
        "table": table,
        "zones": zones,
        "skip_reason": jf.get("skip_reason"),
    }

    peak_raw = jf.get("peak") or {}
    ppar = peak_raw.get("params") or {}
    if ppar:
        out["peak"] = {
            "params": {
                "M_kg_s": _round_if_number(ppar.get("M_kg_s"), 4),
                "LF_m": _round_if_number(ppar.get("LF_m"), 2),
                "DF_m": _round_if_number(ppar.get("DF_m"), 2),
                "Ef_kw_m2": _round_if_number(ppar.get("Ef_kw_m2"), 2),
            },
            "table": _jetfire_table_rows(peak_raw.get("table")),
            "zones": _jetfire_zones_rows(peak_raw.get("zones")),
        }

    def _q_near(table_rows: List[Dict[str, Any]], r_tgt: float) -> Any:
        best_q = None
        best_d = 1e9
        for row in table_rows:
            rm = row.get("r_m")
            if rm is None:
                continue
            try:
                d = abs(float(rm) - r_tgt)
            except (TypeError, ValueError):
                continue
            if d < best_d:
                best_d = d
                best_q = row.get("q_kw_m2")
        return best_q if best_d < 0.51 else None

    out["q_r5_kw_m2"] = _round_if_number(_q_near(table, 5.0), 4)
    out["q_r10_kw_m2"] = _round_if_number(_q_near(table, 10.0), 4)

    # Шаблон ПОУО1 ожидает peak; для прочих сценариев дублируем установившийся режим
    if not out.get("peak") or not (out.get("peak") or {}).get("params"):
        out["peak"] = {
            "params": dict(out["params"]),
            "table": list(out["table"]),
            "zones": list(out["zones"]),
        }

    return out


def _build_fireball_block(results: Dict[str, Any]) -> Dict[str, Any]:
    """
    Подготовка блока огненного шара.

    Из results["fireball"] читаем:
      params  — Ds_m, H_m, ts_s, m_kg, Ef_kw_m2
      table[] — r_m, tau, Fq, q_kw_m2, Pr, prob
      zones[] — q_thr_kw_m2, r_m
    """
    fb = (results or {}).get("fireball", {}) or {}
    params = fb.get("params", {}) or {}

    table = []
    for row in (fb.get("table") or []):
        table.append({
            "r_m":     _round_if_number(row.get("r_m")),
            "tau":     _round_if_number(row.get("tau"), 6),
            "Fq":      _round_if_number(row.get("Fq"), 6),
            "q_kw_m2": _round_if_number(row.get("q_kw_m2"), 4),
            "Pr":      _round_if_number(row.get("Pr"), 4),
            "prob":    _round_if_number(row.get("prob"), 2),
        })

    zones = []
    for z in (fb.get("zones") or []):
        zones.append({
            "q_thr_kw_m2": _round_if_number(z.get("q_thr_kw_m2"), 2),
            "r_m": _pretty_value(z.get("r_m")),
        })

    return {
        "params": {k: _round_if_number(v) for k, v in params.items()},
        "table": table,
        "zones": zones,
        "skip_reason": fb.get("skip_reason"),
    }


def _build_tvs_block(results: Dict[str, Any]) -> Dict[str, Any]:
    """
    Подготовка блока ТВС.

    Важно:
    шаблон использует не только основные поля таблицы,
    но и probit-поля:
      - Pr_people
      - prob_people
      - Pr_full
      - prob_full
      - Pr_heavy
      - prob_heavy
    """
    tvs = (results or {}).get("tvs_explosion", {}) or {}

    inputs = tvs.get("inputs", {}) or {}
    intermediate = tvs.get("intermediate", {}) or {}
    res = tvs.get("results", {}) or {}
    table = tvs.get("table", []) or []

    tvs_table = []
    for row in table:
        tvs_table.append({
            "r_m": _round_if_number(row.get("r_m")),
            "Rx": _round_if_number(row.get("Rx"), 6),
            "Px": _round_if_number(row.get("Px"), 6),
            "Ix": _round_if_number(row.get("Ix"), 6),
            "deltaP_Pa": _round_if_number(row.get("deltaP_Pa"), 4),
            "deltaP_kPa": _round_if_number(row.get("deltaP_kPa"), 4),
            "Iplus_Pa_s": _round_if_number(row.get("Iplus_Pa_s"), 6),

            "Pr_people": _round_if_number(row.get("Pr_people"), 4),
            "prob_people": _round_if_number(row.get("prob_people"), 2),

            "Pr_full": _round_if_number(row.get("Pr_full"), 4),
            "prob_full": _round_if_number(row.get("prob_full"), 2),

            "Pr_heavy": _round_if_number(row.get("Pr_heavy"), 4),
            "prob_heavy": _round_if_number(row.get("prob_heavy"), 2),
        })

    max_delta_p_kpa = None
    max_delta_r_m = None
    max_iplus_pa_s = None
    if tvs_table:
        max_row = max(tvs_table, key=lambda r: r.get("deltaP_Pa") or 0.0)
        max_delta_p_kpa = _round_if_number((max_row.get("deltaP_Pa") or 0.0) / 1000.0, 4)
        max_delta_r_m = max_row.get("r_m")
        max_iplus_row = max(tvs_table, key=lambda r: r.get("Iplus_Pa_s") or 0.0)
        max_iplus_pa_s = _round_if_number(max_iplus_row.get("Iplus_Pa_s"), 2)

    # Радиус зоны «лёгкий вред здоровью» (ΔP ≥ people_light_kPa = 12 кПа)
    _people_light_r = res.get("zones_people", {}).get("people_light_kPa")
    zone_light_harm_r_m = (
        _round_if_number(float(_people_light_r), 1)
        if isinstance(_people_light_r, (int, float))
        else None
    )

    zones_buildings = _pad_building_zones(
        _pretty_building_zones(res.get("zones_buildings", {})),
        5,
    )

    return {
        "inputs": _to_dict(inputs),
        "intermediate": {
            k: _round_if_number(v, 6)
            for k, v in intermediate.items()
            if not isinstance(v, (list, dict, tuple))
        },
        "results": _to_dict(res),
        "logs": tvs.get("logs", []) or [],
        "table": tvs_table,

        "max_delta_p_kpa": max_delta_p_kpa,
        "max_delta_r_m": max_delta_r_m,
        "max_iplus_pa_s": max_iplus_pa_s,
        "zone_light_harm_r_m": zone_light_harm_r_m,

        "zones_glass": _pretty_dict(res.get("zones_glass", {})),
        "zones_people": _pretty_dict(res.get("zones_people", {})),
        "zones_buildings": zones_buildings,

        "flame_speed_m_s": _round_if_number(tvs.get("flame_speed_m_s"), 4),
        "pr4_r0": _round_if_number(tvs.get("pr4_r0"), 4),

        "skip_reason": tvs.get("skip_reason"),
    }


def _build_pool_fire_block(results: Dict[str, Any]) -> Dict[str, Any]:
    """
    Подготовка блока пожара пролива (diesel tank park).
    """
    pf = (results or {}).get("pool_fire", {}) or {}
    params = pf.get("params", {}) or {}

    zones = []
    for z in (pf.get("zones") or []):
        zones.append({
            "q_thr_kw_m2": _round_if_number(z.get("q_thr_kw_m2"), 2),
            "r_m": _pretty_value(z.get("r_m")),
        })

    return {
        "params": {
            "area_m2":    _round_if_number(params.get("area_m2"), 1),
            "d_eff_m":    _round_if_number(params.get("d_eff_m"), 2),
            "H_flame_m":  _round_if_number(params.get("H_flame_m"), 2),
            "Ef_kw_m2":   _round_if_number(params.get("Ef_kw_m2"), 1),
        },
        "zones": zones,
        "skip_reason": pf.get("skip_reason"),
    }


def _build_tank_park_block(results: Dict[str, Any]) -> Dict[str, Any]:
    """
    Сводный блок резервуарного парка для Word-шаблона.

    Агрегирует промежуточные данные из results["tank_park"]["intermediate"]
    и оценку числа людей из results["people_exposure"].
    """
    tp = (results or {}).get("tank_park", {}) or {}
    inter = tp.get("intermediate", {}) or {}
    tp_in = tp.get("inputs", {}) or {}
    meta = tp_in.get("meta", {}) or {}
    fuel_id = str(meta.get("fuel_id", "") or "")

    tank_in = tp_in.get("tank", {}) or {}
    lpg_in = tp_in.get("lpg", {}) or {}
    site_in = tp_in.get("site", {}) or {}

    pe = (results or {}).get("people_exposure", {}) or {}

    def _pe_zones(key: str):
        return [
            {
                "q_thr_kw_m2": _round_if_number(z.get("q_thr_kw_m2"), 2),
                "r_m":         _pretty_value(z.get("r_m")),
                "area_ha":     _round_if_number(z.get("area_ha"), 3),
                "n_people":    _pretty_value(z.get("n_people"), 1),
            }
            for z in (pe.get(key) or [])
        ]

    vol = tank_in.get("volume_m3")
    if vol is None:
        vol = inter.get("volume_m3")
    cnt = tank_in.get("count")
    if cnt is None:
        cnt = inter.get("count")
    try:
        fill = float(tank_in.get("fill_fraction", 0.8) or 0.8)
    except (TypeError, ValueError):
        fill = 0.8

    V_liquid_m3 = None
    if vol is not None and cnt is not None:
        try:
            V_liquid_m3 = float(vol) * int(cnt) * fill
        except (TypeError, ValueError):
            V_liquid_m3 = None

    try:
        exp_f = float(lpg_in.get("expansion_factor", 250) or 250)
    except (TypeError, ValueError):
        exp_f = 250.0
    try:
        rho_v = float(lpg_in.get("rho_vapor_kg_m3", 1.83) or 1.83)
    except (TypeError, ValueError):
        rho_v = 1.83

    V_gvs_m3 = (V_liquid_m3 * exp_f) if V_liquid_m3 is not None else None
    m_gvs_kg = (V_gvs_m3 * rho_v) if V_gvs_m3 is not None else None

    nozzle_r = lpg_in.get("nozzle_radius_m")
    A_hol = None
    if nozzle_r is not None:
        try:
            nr = float(nozzle_r)
            if nr > 0:
                A_hol = round(math.pi * nr * nr, 6)
        except (TypeError, ValueError):
            pass

    P_v = lpg_in.get("P_vessel_Pa")
    P_c = lpg_in.get("P_crit_Pa")
    PR1 = PR2 = None
    P_vessel_MPa = P_crit_MPa = P_atm_MPa = None
    try:
        P_atm = float((tp_in.get("env") or {}).get("P0_Pa", 101_325) or 101_325)
        P_atm_MPa = round(P_atm / 1e6, 3)
    except (TypeError, ValueError):
        P_atm = 101_325.0
        P_atm_MPa = round(P_atm / 1e6, 3)
    try:
        if P_v is not None and P_c is not None and float(P_c) != 0:
            PR1 = round(float(P_v) / float(P_c), 9)
    except (TypeError, ValueError):
        pass
    try:
        if P_c is not None and float(P_c) != 0:
            PR2 = round(P_atm / float(P_c), 9)
    except (TypeError, ValueError):
        pass
    try:
        if P_v is not None:
            P_vessel_MPa = round(float(P_v) / 1e6, 2)
    except (TypeError, ValueError):
        pass
    try:
        if P_c is not None:
            P_crit_MPa = round(float(P_c) / 1e6, 2)
    except (TypeError, ValueError):
        pass

    wind = _wind_zones_from_m_dot(inter.get("m_dot_kg_s"))

    jf_raw = (results or {}).get("jet_fire") or {}
    jf_tab = jf_raw.get("table") or []

    def _site_q(key: str) -> float | None:
        d = site_in.get(key)
        return _interp_q_kw_m2(jf_tab, d)

    # Радиус зоны 7 кВт/м² (санитарные потери) из таблицы зон факела
    zone_sanitary_r_m = None
    for z in (jf_raw.get("zones") or []):
        try:
            if abs(float(z.get("q_thr_kw_m2", 0)) - 7.0) < 0.06:
                zr = z.get("r_m")
                zone_sanitary_r_m = _round_if_number(float(zr), 1) if zr is not None else None
                break
        except (TypeError, ValueError):
            continue

    jf_params = jf_raw.get("params") or {}
    direct_flame_r_m = _round_if_number(jf_params.get("DF_m"), 1)

    return {
        "fuel_id":        fuel_id,
        "volume_m3":      _round_if_number(inter.get("volume_m3"), 1),
        "count":          inter.get("count"),
        "m_total_kg":     _round_if_number(inter.get("m_total_kg"), 1),
        # diesel
        "W_evap_kg_m2_s": _round_if_number(inter.get("W_evap_kg_m2_s"), 6),
        "m_evap_kg":      _round_if_number(inter.get("m_evap_kg"), 2),
        "m_cloud_kg":     _round_if_number(inter.get("m_cloud_kg"), 2),
        "E_J":            _round_if_number(inter.get("E_J"), 2),
        # lpg
        "m_flash_kg":     _round_if_number(inter.get("m_flash_kg"), 2),
        "m_pool_evap_kg": _round_if_number(inter.get("m_pool_evap_kg"), 2),
        "m_dot_kg_s":     _round_if_number(inter.get("m_dot_kg_s"), 4),
        "m_dot_peak_kg_s": _round_if_number(inter.get("m_dot_peak_kg_s"), 2),
        "fill_fraction":  _round_if_number(fill, 2),
        "V_liquid_m3":    _round_if_number(V_liquid_m3, 2),
        "V_gvs_m3":       _round_if_number(V_gvs_m3, 1),
        "m_gvs_kg":       _round_if_number(m_gvs_kg, 1),
        "nozzle_area_m2": A_hol,
        "PR1":            PR1,
        "PR2":            PR2,
        "P_vessel_MPa":   P_vessel_MPa,
        "P_crit_MPa":     P_crit_MPa,
        "P_atm_MPa":      P_atm_MPa,
        "L_wind1_m":      wind["L_wind1_m"],
        "L_wind3_m":      wind["L_wind3_m"],
        "r0_wind1_m":     wind["r0_wind1_m"],
        "r0_wind3_m":     wind["r0_wind3_m"],
        "q_kpp_kw_m2":    _round_if_number(_site_q("dist_kpp_m"), 2),
        "q_sklad_kw_m2":  _round_if_number(_site_q("dist_sklad_m"), 2),
        "q_kotelnaya_kw_m2": _round_if_number(_site_q("dist_kotelnaya_m"), 2),
        "zone_sanitary_r_m": zone_sanitary_r_m,
        "direct_flame_r_m": direct_flame_r_m,
        # люди
        "people_density_per_ha": pe.get("density_per_ha", 0),
        "people_jet_fire":  _pe_zones("jet_fire"),
        "people_pool_fire": _pe_zones("pool_fire"),
        "people_fireball":  _pe_zones("fireball"),
    }


def build_context(project, doc: DocxTemplate | None = None) -> Dict[str, Any]:
    """
    Собирает итоговый контекст для шаблона Word.

    Здесь:
    - читаем project и все его pouos;
    - приводим всё к обычным dict/list;
    - подготавливаем блоки release / tvs / jet_fire / fireball;
    - подготавливаем список труб в удобном для шаблона виде;
    - добавляем пути и InlineImage для графиков.
    """
    ctx: Dict[str, Any] = {
        "project": _dot_for_jinja({
            "name": getattr(project, "name", ""),
            "object_name": getattr(project, "object_name", ""),
            "address": getattr(project, "address", ""),
        }),
        "pouos": [],
    }

    pouos = getattr(project, "pouos", []) or []

    for p in pouos:
        fuel = get_fuel(getattr(p, "fuel_id", ""))

        is_indoor = bool(getattr(p, "is_indoor", False))
        space_title = "Помещение" if is_indoor else "Открытая площадка"

        raw_inputs = _to_dict(getattr(p, "inputs", {}) or {})
        raw_results = _to_dict(getattr(p, "results", {}) or {})
        raw_pipes = _to_dict(getattr(p, "pipes", []) or [])

        merged_inputs = dict(raw_inputs)
        _tpr = raw_results.get("tank_park") or {}
        _tpi = _tpr.get("inputs") or {}
        if _tpi.get("lpg"):
            merged_inputs["lpg"] = _to_dict(_tpi["lpg"])
        if _tpi.get("site"):
            merged_inputs["site"] = _to_dict(_tpi["site"])
        if _tpi.get("tank"):
            _ut = dict(merged_inputs.get("tank") or {})
            for _k, _v in _to_dict(_tpi["tank"]).items():
                if _v is not None:
                    _ut[_k] = _v
            merged_inputs["tank"] = _ut

        # ---------------- Подготовка труб для шаблона ----------------
        # В шаблоне используется pipe.d_inner_m,
        # но в исходных объектах чаще всего есть только diameter_mm.
        # Поэтому создаём подготовленный список труб.
        prepared_pipes = []
        for pipe in raw_pipes:
            diameter_mm = pipe.get("diameter_mm")
            d_inner_m = None

            if diameter_mm is not None:
                try:
                    d_inner_m = round(float(diameter_mm) / 1000.0, 3)
                except (TypeError, ValueError):
                    d_inner_m = None

            prepared_pipes.append({
                **pipe,
                "d_inner_m": d_inner_m,
                "length_m": pipe.get("length_m"),
            })

        # ---------------- Подготовка расчётных блоков ----------------
        release_block = _build_release_block(raw_results)
        jetfire_block = _build_jetfire_block(raw_results)
        fireball_block = _build_fireball_block(raw_results)
        tvs_block = _build_tvs_block(raw_results)
        pool_fire_block = _build_pool_fire_block(raw_results)
        tank_park_block = _build_tank_park_block(raw_results)

        code = getattr(p, "code", "")
        charts_dir = os.path.join("out", "charts")

        tvs_dp_path = os.path.join(charts_dir, f"tvs_dp_{code}.png")
        tvs_imp_path = os.path.join(charts_dir, f"tvs_imp_{code}.png")
        jetfire_path = os.path.join(charts_dir, f"jetfire_{code}.png")
        fireball_path = os.path.join(charts_dir, f"fireball_{code}.png")

        p_dict: Dict[str, Any] = {
            "code": code,
            "title": getattr(p, "title", ""),
            "is_indoor": is_indoor,
            "space_title": space_title,

            "fuel_id": getattr(fuel, "id", getattr(p, "fuel_id", "")),
            "fuel_title": getattr(fuel, "title", ""),
            "eud0_j_per_kg": getattr(fuel, "eud0_j_per_kg", 0.0),

            # Сырой слой (для tank_park дополняем lpg/site/fill из снимка inputs расчёта)
            "inputs": merged_inputs,
            "results": raw_results,

            # Подготовленные трубы для таблицы Word
            "pipes": prepared_pipes,

            # Подготовленные блоки расчёта
            "release": release_block,
            "jet_fire": jetfire_block,
            "fireball": fireball_block,
            "tvs": tvs_block,
            "pool_fire": pool_fire_block,
            "tank_park": tank_park_block,

            # Пути к графикам
            "tvs_dp_chart_path": tvs_dp_path,
            "tvs_imp_chart_path": tvs_imp_path,
            "jetfire_chart_path": jetfire_path,
            "fireball_chart_path": fireball_path,

            # Флаги для шаблона
            "has_release":   bool(release_block   and not release_block.get("skip_reason")),
            "has_jet_fire":  bool(jetfire_block    and not jetfire_block.get("skip_reason")),
            "has_fireball":  bool(fireball_block   and not fireball_block.get("skip_reason")),
            "has_tvs":       bool(tvs_block        and not tvs_block.get("skip_reason")),
            "has_pool_fire": bool(pool_fire_block  and not pool_fire_block.get("skip_reason")
                                  and pool_fire_block.get("zones")),
            "has_tank_park": bool(raw_results.get("tank_park")),
            "is_tank_park":  bool(raw_results.get("tank_park")),
            "has_error":     bool(raw_results.get("error")),
            "error_text":    raw_results.get("error"),
        }

        # Если передан doc — подцепляем картинки как InlineImage
        if doc is not None:
            p_dict["tvs_dp_chart_img"] = _safe_inline_image(doc, tvs_dp_path, width_mm=150)
            p_dict["tvs_imp_chart_img"] = _safe_inline_image(doc, tvs_imp_path, width_mm=150)
            p_dict["jetfire_chart_img"] = _safe_inline_image(doc, jetfire_path, width_mm=150)
            p_dict["fireball_chart_img"] = _safe_inline_image(doc, fireball_path, width_mm=150)

        ctx["pouos"].append(_dot_for_jinja(p_dict))

    return ctx


def render_report(template_path: str, output_path: str, project) -> None:
    """
    Рендерит Word-файл по шаблону docxtpl.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Не найден шаблон: {template_path}")

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    logger.debug("Загружаем шаблон: %s", template_path)
    doc = DocxTemplate(template_path)

    logger.debug("Строим контекст для docxtpl")
    ctx = build_context(project, doc=doc)

    logger.debug("Рендерим шаблон")
    doc.render(ctx)

    logger.debug("Сохраняем: %s", output_path)
    doc.save(output_path)

    logger.debug("Отчёт сохранён")


def render_report_for_project(
    project: Any,
    output_path: str,
    templates_dir: str | Path | None = None,
) -> str:
    """
    Подбирает шаблон по **единственному** типу ПОУО в проекте, рендерит один .docx.

    Если в проекте несколько разных ``code``, будет ошибка — используйте
    :func:`render_full_report`.

    Возвращает имя использованного файла шаблона.
    """
    template_path, filename = resolve_word_template_path(project, templates_dir)
    render_report(str(template_path), output_path, project)
    return filename


def group_pouos_by_code(project: Any) -> Tuple[Dict[str, List[Any]], List[str]]:
    """
    Группировка ПОУО по полю code.

    Возвращает (groups, order), где order — порядок первого появления каждого code
    в project.pouos (как добавлял пользователь).
    """
    pouos = getattr(project, "pouos", None) or []
    groups: Dict[str, List[Any]] = {}
    order: List[str] = []

    for p in pouos:
        code = getattr(p, "code", "") or ""
        if code not in groups:
            groups[code] = []
            order.append(code)
        groups[code].append(p)

    return groups, order


def render_reports_per_scenario(
    project: Any,
    output_dir: str | Path,
    templates_dir: str | Path | None = None,
) -> Dict[str, str]:
    """
    Для каждого встречающегося code — отдельный файл ``report_{code}.docx`` в output_dir.
    Порядок генерации совпадает с порядком первого появления code в project.pouos.

    Проект для каждого файла — deepcopy исходного с подсписком pouos только этого code.

    Возвращает словарь code → абсолютный путь (в Python 3.7+ порядок ключей = порядок обхода).
    """
    groups, order = group_pouos_by_code(project)
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    base_templates = Path(templates_dir) if templates_dir else default_word_templates_dir()

    saved: Dict[str, str] = {}
    for code in order:
        sub = copy.deepcopy(project)
        sub.pouos = [p for p in sub.pouos if getattr(p, "code", "") == code]
        template_path, _ = resolve_word_template_path(sub, base_templates)
        out_path = out_dir / f"report_{code}.docx"
        render_report(str(template_path), str(out_path), sub)
        saved[code] = str(out_path.resolve())
    return saved


def _remove_default_empty_paragraph(doc: Any) -> None:
    """Убирает начальный пустой абзац у только что созданного Document()."""
    paras = doc.paragraphs
    if not paras:
        return
    first = paras[0]
    if first.text.strip() == "" and len(first.runs) == 0:
        el = first._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def add_title_page(doc: Any, project: Any) -> None:
    """Титульный лист в стиле ГОСТ (организация, тема, объект, исполнитель)."""
    uni = doc.add_paragraph()
    uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = uni.add_run(TITLE_PAGE_UNIVERSITY)
    r.bold = True
    r.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(TITLE_PAGE_UNIVERSITY_SHORT)
    r2.bold = True
    r2.font.size = Pt(14)

    fac = doc.add_paragraph()
    fac.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fac.add_run(TITLE_PAGE_FACULTY)

    dep = doc.add_paragraph()
    dep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dep.add_run(TITLE_PAGE_DEPARTMENT)

    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Паспорт безопасности объекта ТЭК")
    tr.bold = True
    tr.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Отчёт по расчётам последствий аварий")

    doc.add_paragraph()
    doc.add_paragraph()

    obj_name = getattr(project, "object_name", "") or "—"
    addr = getattr(project, "address", "") or "—"
    doc.add_paragraph(f"Объект: {obj_name}")
    doc.add_paragraph(f"Адрес: {addr}")

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    year = datetime.now().year
    sign.add_run(
        f"Выполнил: {TITLE_PAGE_AUTHOR}\n"
        f"Группа: {TITLE_PAGE_GROUP}\n"
        f"Год: {year}"
    )

    doc.add_page_break()


def add_table_of_contents(doc: Any) -> None:
    """Поле оглавления Word (уровни 1–3); после открытия в Word — обновить поля."""
    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = head.add_run("Содержание")
    hr.bold = True
    hr.font.size = Pt(14)

    doc.add_paragraph()

    p_toc = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u ')
    p_toc._p.append(fld)

    hint = doc.add_paragraph()
    hint_run = hint.add_run(
        "(В Microsoft Word: выделите оглавление → ПКМ → «Обновить поле».)"
    )
    hint_run.italic = True
    hint_run.font.size = Pt(9)

    doc.add_page_break()


def export_to_pdf(docx_path: str | Path, pdf_path: str | Path) -> None:
    """Экспорт .docx в .pdf через Word (Windows). Требуется пакет docx2pdf."""
    try:
        from docx2pdf import convert
    except ImportError as e:
        raise ImportError(
            "Для экспорта в PDF установите зависимость: pip install docx2pdf"
        ) from e
    dx = str(Path(docx_path).resolve())
    pdf = str(Path(pdf_path).resolve())
    Path(pdf).parent.mkdir(parents=True, exist_ok=True)
    convert(dx, pdf)


def _apply_landscape_a4_all_sections(doc: Any) -> None:
    """Все секции документа — альбомная A4 (297×210 мм)."""
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)


def _normal_style(doc: Any):
    try:
        return doc.styles["Normal"]
    except (KeyError, ValueError):
        return None


def _strip_paragraph_background(doc: Any, paragraph: Any) -> None:
    """Highlight, w:rPr/w:shd, w:pPr/w:shd, сброс стиля на Normal (уменьшает заливку из стилей)."""
    for run in paragraph.runs:
        try:
            run.font.highlight_color = None
        except (AttributeError, TypeError):
            pass
        r_el = run._element
        r_pr = r_el.find(qn("w:rPr"))
        if r_pr is not None:
            shd = r_pr.find(qn("w:shd"))
            if shd is not None:
                r_pr.remove(shd)

    p_el = paragraph._element
    p_pr = p_el.find(qn("w:pPr"))
    if p_pr is not None:
        shd = p_pr.find(qn("w:shd"))
        if shd is not None:
            p_pr.remove(shd)

    norm = _normal_style(doc)
    if norm is not None:
        try:
            paragraph.style = norm
        except (ValueError, KeyError):
            pass


def _remove_cell_tc_shading(cell: Any) -> None:
    tc = cell._element
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is not None:
        shd = tc_pr.find(qn("w:shd"))
        if shd is not None:
            tc_pr.remove(shd)


def _remove_background_from_table(doc: Any, table: Any) -> None:
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is not None:
        shd = tbl_pr.find(qn("w:shd"))
        if shd is not None:
            tbl_pr.remove(shd)

    for row in table.rows:
        tr_pr = row._tr.find(qn("w:trPr"))
        if tr_pr is not None:
            shd = tr_pr.find(qn("w:shd"))
            if shd is not None:
                tr_pr.remove(shd)
        for cell in row.cells:
            _remove_cell_tc_shading(cell)
            for paragraph in cell.paragraphs:
                _strip_paragraph_background(doc, paragraph)
            for nested in cell.tables:
                _remove_background_from_table(doc, nested)


def remove_background(doc: Any) -> None:
    """
    Убирает заливку: highlight, w:shd на run/абзаце/ячейке, сбрасывает стиль абзацев на Normal.

    Обходит также вложенные таблицы. Полностью убрать заливку из определений стилей в .docx
    без правки styles.xml нельзя; сброс на Normal снимает применение многих стилей с фоном.
    """
    for paragraph in doc.paragraphs:
        _strip_paragraph_background(doc, paragraph)
    for table in doc.tables:
        _remove_background_from_table(doc, table)


def merge_word_files(
    files: Sequence[str | Path],
    output_path: str | Path,
    project: Any,
    scenario_codes: Sequence[str],
) -> None:
    """
    Собирает итоговый документ: пустой ``Document()``, титульник, оглавление,
    затем сценарии с заголовками «Сценарий N — CODE» и телом каждого фрагмента.

    Объединение фрагментов через ``docxcompose.Composer`` (relationships, media, стили).
    Прямой перенос ``element.body`` не используется.

    Между сценариями — разрыв страницы. Порядок ``files`` и ``scenario_codes`` должен совпадать.
    """
    try:
        from docxcompose.composer import Composer
    except ImportError as e:
        raise ImportError(
            "Для сборки итогового Word установите: pip install docxcompose"
        ) from e

    paths = [str(p) for p in files]
    codes = list(scenario_codes)
    if not paths:
        raise ValueError("merge_word_files: пустой список файлов")
    if len(paths) != len(codes):
        raise ValueError(
            "merge_word_files: число файлов и кодов сценариев должно совпадать"
        )

    base = _DocxDocument()
    _remove_default_empty_paragraph(base)
    _apply_landscape_a4_all_sections(base)

    add_title_page(base, project)
    add_table_of_contents(base)

    composer = Composer(base)

    for i, (path, code) in enumerate(zip(paths, codes)):
        if i > 0:
            base.add_page_break()
        base.add_heading(f"Сценарий {i + 1} — {code}", level=1)
        composer.append(_DocxDocument(path))

    _apply_landscape_a4_all_sections(base)
    remove_background(base)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(out))


def render_full_report(
    project: Any,
    output_path: str | Path,
    templates_dir: str | Path | None = None,
    *,
    export_pdf: bool = False,
) -> None:
    """
    Временные ``report_{code}.docx`` по порядку добавления ПОУО → один итоговый файл
    (титульник, оглавление, пронумерованные сценарии).

    Временная папка создаётся через ``tempfile`` и удаляется после merge.

    При ``export_pdf=True`` рядом сохраняется файл с расширением ``.pdf`` (нужен docx2pdf и Word).
    """
    out = Path(output_path)
    with tempfile.TemporaryDirectory(prefix="tek_passport_word_") as tmp:
        per_code = render_reports_per_scenario(project, tmp, templates_dir)
        order = list(per_code.keys())
        files_in_order = list(per_code.values())
        if not files_in_order:
            raise ValueError("В проекте нет ПОУО — нечего включать в отчёт.")
        merge_word_files(files_in_order, out, project, order)

    if export_pdf:
        pdf_path = out.with_suffix(".pdf")
        export_to_pdf(out, pdf_path)